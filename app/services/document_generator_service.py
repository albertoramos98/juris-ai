import io
import os
from docx import Document

def generate_docx_from_template(template_path: str, replacements: dict) -> io.BytesIO:
    """
    Abre um modelo .docx e substitui placeholders preservando RIGOROSAMENTE
    a formatação original (negrito, itálico, fontes, etc).
    """
    if not os.path.exists(template_path):
        doc = Document()
        doc.add_heading("Erro: Modelo não localizado", 0)
        doc.add_paragraph(f"Caminho: {template_path}")
    else:
        doc = Document(template_path)

    def smart_replace(paragraphs, replacements):
        for p in paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    # Percorre os 'runs' (fragmentos de formatação) do parágrafo
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
                    
                    # Fallback para o parágrafo todo se ainda sobrar a chave
                    if key in p.text: 
                        p.text = p.text.replace(key, str(value))

    # 1. Substituir no corpo
    smart_replace(doc.paragraphs, replacements)

    # 2. Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                smart_replace(cell.paragraphs, replacements)

    # 3. Substituir em Cabeçalhos e Rodapés
    for section in doc.sections:
        smart_replace(section.header.paragraphs, replacements)
        smart_replace(section.footer.paragraphs, replacements)
        if section.different_first_page_header_footer:
            smart_replace(section.first_page_header.paragraphs, replacements)
            smart_replace(section.first_page_footer.paragraphs, replacements)

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

def generate_docx_from_text(title: str, full_text: str) -> io.BytesIO:
    """
    Gera um DOCX básico do zero. 
    Mantido para compatibilidade com o endpoint rag_export_docx original.
    """
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    
    if full_text:
        doc.add_paragraph(full_text)
        
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target
